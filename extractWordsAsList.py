# extract utf-8 codes from file textToConvert.docx and make a list extractedWords
import pandas as pd
import aspose.words as aw
import docx

words = []

# # # convert to utf-8
# file=open('inputTextToConvertHere.txt', mode='rt',encoding='utf-8')
# newfile=open('textToConvert.txt', mode='a')
# for eachLine in file:
#     eachLine=eachLine.strip()
#     newLine=str(eachLine.encode("utf-8"))
#     newfile.write(newLine[2:]+'\n')

# #replace last ' with back space and \ with |
# with open('textToConvert.txt','r') as file:
#     data=file.read()
#     data=data.replace("\\",'|')
        
#     data=data.replace("'",'\b')
#     # data=data.replace("\n\n",'\n')

# with open('textToConvert.txt','w') as file:
#     file.write(data)

# #convert textToConvert extension to docx
# #this code is not running on ubuntu for some reason
# #so i manually created textToConvert.docx
# input=aw.Document("textToConvert.txt")
# input.save("textToConvert.docx")


def getText(filename):
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
        words.append(para.text)
    return " ".join(fullText)

getText(r"textToConvert.docx")

count= len(words)
print("Total=", count)
extractedWords=[]
for i in range (0, count):
    curcheck= words[i].split(" ")
    extractedWords.append(curcheck)
    
for j in extractedWords:
    print (j)

with open("extractedWords.txt","w", encoding="utf-8") as output:
    for listitem in extractedWords:
        output.write('%s\n' %listitem)
    